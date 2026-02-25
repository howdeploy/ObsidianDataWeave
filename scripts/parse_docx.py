"""
parse_docx.py — Convert a .docx file to structured JSON for atomization.

Preserves heading hierarchy (normalized to level 1 = top-level, regardless
of whether the source document uses H1, H3, or any other level as its top),
inline formatting (bold/italic as markdown), list items, tables (as markdown),
and marks images as [image] placeholders.

Usage:
    python3 scripts/parse_docx.py <input.docx>
    python3 scripts/parse_docx.py <input.docx> -o output.json

Output JSON schema:
{
  "source_file": "filename.docx",
  "heading_depth_offset": N,   # e.g. 2 means H3 is top level in source
  "sections": [
    {
      "heading": "Section Title" | null,
      "level": 1,               # normalized: 1 = top, 2 = sub, etc.
      "paragraphs": ["text with **bold** and *italic*..."]
    }
  ]
}
"""

import re
import sys
import json
import argparse
from pathlib import Path

try:
    import docx
except ImportError:
    print(
        "ERROR: python-docx is required. Install with:\n"
        "  pip install python-docx --user",
        file=sys.stderr,
    )
    sys.exit(1)


# ── Heading detection ──────────────────────────────────────────────────────────

def get_heading_level(para) -> int | None:
    """Return the numeric heading level (1-9) for a heading paragraph, or None."""
    m = re.match(r"Heading\s+(\d+)", para.style.name, re.IGNORECASE)
    return int(m.group(1)) if m else None


# ── Inline formatting ──────────────────────────────────────────────────────────

def runs_to_markdown(para) -> str:
    """Convert a paragraph's runs to a string with markdown inline formatting."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            text = f"***{text}***"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts)


# ── List detection ─────────────────────────────────────────────────────────────

_LIST_STYLES = {"list paragraph", "list bullet", "list number", "list continue"}

# Bullet markers that indicate an inline list item even in normal-style paragraphs
_BULLET_RE = re.compile(r"^[\u2022\u2023\u25E6\u2043\u2219\-\*]\s")
_NUMBER_RE = re.compile(r"^\d+[\.\)]\s")


def is_list_item(para) -> bool:
    """Heuristically detect whether a paragraph is a list item."""
    style_name = para.style.name.lower()
    if any(ls in style_name for ls in _LIST_STYLES):
        return True
    text = para.text
    if _BULLET_RE.match(text) or _NUMBER_RE.match(text):
        return True
    return False


def para_to_list_markdown(para) -> str:
    """Format a list paragraph as a markdown list item."""
    text = runs_to_markdown(para).strip()
    # Normalize existing bullet markers to markdown dash
    text = _BULLET_RE.sub("", text)
    text = _NUMBER_RE.sub("", text)
    # Determine indentation level from style name if possible
    # e.g. "List Bullet 2" -> indent level 2
    m = re.search(r"\d+$", para.style.name)
    indent_level = int(m.group()) - 1 if m else 0
    prefix = "  " * indent_level + "- "
    return prefix + text.strip() if text.strip() else ""


# ── Image detection ────────────────────────────────────────────────────────────

def has_inline_image(para) -> bool:
    """Check if paragraph contains an inline image (Drawing XML element)."""
    # python-docx exposes the XML; check for drawing/picture namespace
    xml_str = para._p.xml if hasattr(para._p, "xml") else ""
    return "w:drawing" in xml_str or "pic:pic" in xml_str


# ── Table conversion ───────────────────────────────────────────────────────────

def table_to_markdown(table) -> str:
    """Convert a docx Table to a GitHub-flavored markdown table string."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Header row
    header = rows[0]
    separator = ["---"] * len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


# ── Paragraph position mapping ─────────────────────────────────────────────────

def build_para_position_map(doc) -> dict:
    """Build a map from paragraph XML element id -> sequential position index."""
    return {id(p._p): i for i, p in enumerate(doc.paragraphs)}


def find_table_insert_position(table, para_position_map: dict) -> int:
    """
    Find the position (in the paragraphs list) after which this table should
    be inserted. Returns -1 if the table appears before all paragraphs.
    """
    # The table's parent element is doc.element.body; find the table's position
    # among body children to determine where it falls relative to paragraphs.
    body = table._tbl.getparent()
    if body is None:
        return -1

    last_para_before_table = -1
    found_table = False
    para_idx = -1

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            para_idx = para_position_map.get(id(child), para_idx)
            if not found_table:
                last_para_before_table = para_idx
        elif tag == "tbl" and child is table._tbl:
            found_table = True
            break

    return last_para_before_table


# ── Core parser ────────────────────────────────────────────────────────────────

def parse_docx_to_json(path: str) -> dict:
    """
    Parse a .docx file and return a structured dict.

    Args:
        path: Absolute or relative path to the .docx file.

    Returns:
        dict with keys: source_file, heading_depth_offset, sections.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the path is not a .docx file.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if p.suffix.lower() != ".docx":
        raise ValueError(f"Not a .docx file: {path}")

    doc = docx.Document(str(p))

    # ── Determine heading normalization offset ─────────────────────────────────
    heading_levels = [
        get_heading_level(para)
        for para in doc.paragraphs
        if get_heading_level(para) is not None
    ]
    min_level = min(heading_levels) if heading_levels else 1
    heading_depth_offset = min_level - 1

    # ── Prepare table injection map ────────────────────────────────────────────
    # Build {insert_after_para_index -> [markdown_table_string, ...]}
    para_pos_map = build_para_position_map(doc)
    table_injections: dict[int, list[str]] = {}
    for table in doc.tables:
        insert_at = find_table_insert_position(table, para_pos_map)
        md_table = table_to_markdown(table)
        if md_table:
            table_injections.setdefault(insert_at, []).append(md_table)

    # ── Walk paragraphs ────────────────────────────────────────────────────────
    result: dict = {
        "source_file": p.name,
        "heading_depth_offset": heading_depth_offset,
        "sections": [],
    }
    current_section: dict | None = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Inject any tables that belong after the previous paragraph
        if i in table_injections:
            target_section = current_section
            if target_section is None:
                # Tables before first heading go into a pre-heading section
                target_section = {"heading": None, "level": 0, "paragraphs": []}
                result["sections"].append(target_section)
                current_section = target_section
            for md_table in table_injections[i]:
                target_section["paragraphs"].append(md_table)

        # Skip empty paragraphs
        if not text:
            # Still check for inline images on otherwise-empty paragraphs
            if has_inline_image(para):
                if current_section is None:
                    current_section = {"heading": None, "level": 0, "paragraphs": []}
                    result["sections"].append(current_section)
                current_section["paragraphs"].append("[image]")
            continue

        level = get_heading_level(para)

        if level is not None:
            # New section
            normalized_level = level - heading_depth_offset
            current_section = {
                "heading": text,
                "level": normalized_level,
                "paragraphs": [],
            }
            result["sections"].append(current_section)
        else:
            # Body paragraph — ensure we have a current section
            if current_section is None:
                current_section = {"heading": None, "level": 0, "paragraphs": []}
                result["sections"].append(current_section)

            # Check for inline image
            if has_inline_image(para):
                current_section["paragraphs"].append("[image]")
                continue

            # List item
            if is_list_item(para):
                content = para_to_list_markdown(para)
                if content:
                    current_section["paragraphs"].append(content)
                continue

            # Regular paragraph with inline markdown formatting
            content = runs_to_markdown(para)
            if content.strip():
                current_section["paragraphs"].append(content)

    # Inject any tables that appear after the last paragraph
    if len(doc.paragraphs) in table_injections:
        target_section = current_section
        if target_section is None:
            target_section = {"heading": None, "level": 0, "paragraphs": []}
            result["sections"].append(target_section)
        for md_table in table_injections[len(doc.paragraphs)]:
            target_section["paragraphs"].append(md_table)

    return result


# ── CLI ────────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Parse a .docx file into structured JSON for atomization."
    )
    parser.add_argument("input", help="Path to the .docx file")
    parser.add_argument(
        "-o", "--output", help="Output JSON file path (default: print to stdout)"
    )
    args = parser.parse_args()

    try:
        result = parse_docx_to_json(args.input)
    except FileNotFoundError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: Unexpected error parsing {args.input}: {e}", file=sys.stderr)
        sys.exit(1)

    json_output = json.dumps(result, ensure_ascii=False, indent=2)

    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(json_output, encoding="utf-8")
        print(f"Written to: {output_path}", file=sys.stderr)
    else:
        print(json_output)


if __name__ == "__main__":
    main()
